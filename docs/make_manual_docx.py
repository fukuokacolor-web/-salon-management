# -*- coding: utf-8 -*-
"""
OWNER_SETUP_MANUAL.md をもとに、超初心者向けWord版マニュアル（.docx）を生成する。
python-docx v1.2.0 使用。
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUT_PATH = r"C:\Users\Owner\Documents\salon-repo\docs\OWNER_SETUP_MANUAL.docx"

# カラーパレット
COLOR_H1 = RGBColor(0xD4, 0x68, 0x8A)   # ローズピンク
COLOR_H2 = RGBColor(0x7B, 0x4A, 0x5C)   # ダークピンク
COLOR_WARN = RGBColor(0xE5, 0x39, 0x35) # 赤
COLOR_TIP = RGBColor(0x19, 0x76, 0xD2)  # 青
COLOR_IMPORTANT = RGBColor(0x2E, 0x7D, 0x32) # 緑
COLOR_STEP = RGBColor(0xD4, 0x68, 0x8A)
COLOR_GREY_BG = "F2F2F2"
COLOR_BORDER = "CCCCCC"
COLOR_TABLE_HEAD = "F8E1EA"


FONT_JP = "メイリオ"
FONT_CODE = "Consolas"


def set_cell_shading(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_cell_border(cell, color="CCCCCC", sz="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def set_paragraph_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def set_paragraph_border(paragraph, color="CCCCCC", sz="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:space"), "4")
        b.set(qn("w:color"), color)
        pBdr.append(b)
    pPr.append(pBdr)


def apply_jp_font(run, font_name=FONT_JP):
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def add_run(paragraph, text, *, bold=False, color=None, size=None, font=FONT_JP):
    run = paragraph.add_run(text)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)
    apply_jp_font(run, font)
    return run


def add_heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.style = doc.styles["Heading 1"]
    # clear default runs
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_H1
    apply_jp_font(run)
    # underline border
    set_paragraph_border_bottom(p, color="D4688A", sz="12")
    return p


def set_paragraph_border_bottom(paragraph, color="D4688A", sz="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), sz)
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), color)
    pBdr.append(b)
    pPr.append(pBdr)


def add_heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.style = doc.styles["Heading 2"]
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = COLOR_H2
    apply_jp_font(run)
    return p


def add_heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.style = doc.styles["Heading 3"]
    for r in list(p.runs):
        r.text = ""
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = COLOR_H2
    apply_jp_font(run)
    return p


def add_para(doc, text, *, bold=False, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=bold, color=color, size=size)
    return p


def add_code_block(doc, text):
    """等幅フォント＋グレー背景の擬似コードブロック。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run(text)
    run.font.name = FONT_CODE
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_CODE)
    rFonts.set(qn("w:hAnsi"), FONT_CODE)
    rFonts.set(qn("w:eastAsia"), FONT_JP)
    run.font.size = Pt(9.5)
    set_paragraph_shading(p, COLOR_GREY_BG)
    set_paragraph_border(p, color=COLOR_BORDER, sz="6")
    return p


def add_step_banner(doc, number, title, time_min, difficulty, depends):
    """Step番号バナー + メタ情報テーブル。"""
    # Step番号（大きな装飾）
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"Step {number}")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = COLOR_STEP
    apply_jp_font(run)
    run2 = p.add_run(f"  {title}")
    run2.bold = True
    run2.font.size = Pt(18)
    run2.font.color.rgb = COLOR_H2
    apply_jp_font(run2)
    set_paragraph_border_bottom(p, color="D4688A", sz="18")

    # メタ情報テーブル
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [f"⏱ 所要時間: {time_min}", f"📊 難易度: {difficulty}", f"🔗 前提: {depends}"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        para = cell.paragraphs[0]
        run = para.add_run(h)
        run.font.size = Pt(10)
        run.bold = True
        apply_jp_font(run)
        set_cell_shading(cell, "FDF0F4")
        set_cell_border(cell, color="D4688A", sz="6")
    add_blank(doc)


def add_callout(doc, kind, text):
    """
    kind: 'warn' (⚠️ 注意 / 赤), 'tip' (💡 ヒント / 青), 'important' (🎯 重要 / 緑)
    """
    if kind == "warn":
        icon, color, fill = "⚠️ 注意", COLOR_WARN, "FDECEA"
    elif kind == "tip":
        icon, color, fill = "💡 ヒント", COLOR_TIP, "E8F1FB"
    else:
        icon, color, fill = "🎯 重要", COLOR_IMPORTANT, "E8F5E9"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run1 = p.add_run(f"{icon}  ")
    run1.bold = True
    run1.font.color.rgb = color
    run1.font.size = Pt(11)
    apply_jp_font(run1)
    run2 = p.add_run(text)
    run2.font.color.rgb = color
    run2.font.size = Pt(10.5)
    apply_jp_font(run2)
    set_paragraph_shading(p, fill)
    set_paragraph_border(p, color="D0D0D0", sz="6")


def add_action(doc, kind, text):
    """operate/check/trouble 各サブステップ。"""
    if kind == "operate":
        icon = "👉 操作"
        color = COLOR_H2
    elif kind == "check":
        icon = "✅ 確認"
        color = COLOR_IMPORTANT
    else:
        icon = "⚠ うまくいかない時"
        color = COLOR_WARN
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(icon)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = color
    apply_jp_font(run)
    if text:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(2)
        add_run(p2, text)


def add_bullets(doc, items, numbered=False):
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number" if numbered else "List Bullet")
        p.paragraph_format.space_after = Pt(2)
        # style may reset font; just add a plain run
        for r in list(p.runs):
            r.text = ""
        add_run(p, item)


def add_checklist(doc, items):
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        add_run(p, f"☐  {item}")


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        apply_jp_font(run)
        set_cell_shading(cell, COLOR_TABLE_HEAD)
        set_cell_border(cell, color="CCCCCC", sz="4")
    for ri, row in enumerate(rows, 1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9.5)
            apply_jp_font(run)
            set_cell_border(cell, color="CCCCCC", sz="4")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    add_blank(doc)
    return table


def add_blank(doc):
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


# ---------------- Document Build ----------------

def build():
    doc = Document()

    # ページ設定（A4）
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Normalスタイルの日本語フォント
    style = doc.styles["Normal"]
    style.font.name = FONT_JP
    style.font.size = Pt(10.5)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), FONT_JP)
    rFonts.set(qn("w:ascii"), FONT_JP)
    rFonts.set(qn("w:hAnsi"), FONT_JP)

    # フッター（ページ番号）
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    fr._r.append(fldChar1)
    fr._r.append(instrText)
    fr._r.append(fldChar2)
    apply_jp_font(fr)

    # ===== 表紙 =====
    for _ in range(5):
        add_blank(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🌸 サロン顧客管理システム")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_H1
    apply_jp_font(run)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("セットアップガイド")
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = COLOR_H1
    apply_jp_font(run)

    for _ in range(3):
        add_blank(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("対象：サロンオーナー様")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = COLOR_H2
    apply_jp_font(run)

    add_blank(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("バージョン v77")
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_H2
    apply_jp_font(run)

    add_blank(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("初回読了目安：3時間（実作業2時間＋読む時間1時間）")
    run.font.size = Pt(12)
    run.font.color.rgb = COLOR_H2
    apply_jp_font(run)

    for _ in range(6):
        add_blank(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("最終更新：2026-04-20")
    run.font.size = Pt(10)
    apply_jp_font(run)

    add_page_break(doc)

    # ===== 目次 =====
    add_heading1(doc, "📖 目次")
    toc_items = [
        ("1. はじめに", "このシステムで何ができるか／準備するもの"),
        ("2. 事前準備", "Googleアカウント・PC・Chromeの準備"),
        ("Step 1", "Googleスプレッドシートをコピーする"),
        ("Step 2", "LINE公式アカウントを作る"),
        ("Step 3", "GAS（プログラム）を自分用に設定する"),
        ("Step 4", "ウェブアプリとして公開する（デプロイ）"),
        ("Step 5", "LINEのWebhook設定"),
        ("Step 6", "初期化スクリプトを実行する"),
        ("Step 7", "リッチメニューを設定する"),
        ("Step 8", "管理画面にログインしてみる"),
        ("Step 9", "テスト予約をしてみる"),
        ("Step 10", "定期トリガーをセットする"),
        ("日常運用ガイド", "毎日・週1・月1のチェック"),
        ("FAQ", "こんなとき どうする"),
        ("困ったら連絡先", "問い合わせテンプレ"),
        ("用語集", "GAS/Webhook/デプロイ などの初心者向け定義"),
    ]
    add_table(doc, ["章", "内容"], toc_items, col_widths=[4.0, 12.0])

    add_page_break(doc)

    # ===== 1. はじめに =====
    add_heading1(doc, "1. はじめに")
    add_para(doc, "ようこそ！このマニュアルは、あなたのサロンに「サロン顧客管理システム」を導入するための完全ガイドです。")
    add_callout(doc, "tip",
        "このマニュアルは、パソコン操作が苦手な方でも読めるように書かれています。"
        "専門用語が出てきたら、そのつど「◯◯のこと」と言い換えています。"
        "分からないことがあっても焦らず、最後の「困ったら連絡先」まで進めば大丈夫です。"
        "所要時間の目安は合計2〜3時間。1日で全部やらなくてもOKです。")

    add_heading2(doc, "🎉 このシステムで何ができるの？")
    add_para(doc, "このシステムを導入すると、あなたのサロンで以下のことが自動でできるようになります。")
    add_bullets(doc, [
        "📱 LINEからお客様が24時間いつでも予約できる（カレンダーをタップするだけ）",
        "💾 お客様情報・来店履歴・ポイントが自動で蓄積される",
        "📊 ダッシュボードで本日の予約・売上・コース残回数が一目でわかる",
        "💌 前日に自動でリマインドLINEが届く（ドタキャン防止）",
        "🔐 お客様ご自身がポイント残高や予約を確認できる",
        "🗄 毎週日曜に自動でバックアップが取られる",
    ])

    add_heading2(doc, "📋 準備するもの一覧（全部無料でOK）")
    add_table(doc,
        ["#", "必要なもの", "補足"],
        [
            ["1", "Googleアカウント", "Gmail を使っている方ならOK"],
            ["2", "LINE公式アカウント", "無料プランでOK（月200通まで無料）"],
            ["3", "PC（Windows / Mac）", "スマホだけでは設定できません"],
            ["4", "開発者からの引き渡しパック",
             "①スプレッドシート共有リンク ②GAS共有リンク ③本マニュアル ④管理画面URL ⑤運用形態（A/B）"],
        ],
        col_widths=[1.2, 4.5, 10.0])
    add_callout(doc, "warn",
        "開発者から「引き渡しパック」がまだ届いていない場合は、先に開発者に連絡してください（最終章「困ったら連絡先」参照）。")

    add_page_break(doc)

    # ===== 2. 事前準備 =====
    add_heading1(doc, "2. 事前準備")
    add_heading2(doc, "✅ Googleアカウントの準備")
    add_para(doc, "すでに Gmail を使っている方は、そのアカウントで問題ありません。")
    add_bullets(doc, [
        "サロン専用の Gmail を作ることを強く推奨します（例: yoursalon@gmail.com）。",
        "プライベートの Gmail と混ぜないほうが、将来スタッフに引き継ぐときに楽です。",
    ])
    add_heading2(doc, "✅ LINE公式アカウントの準備")
    add_para(doc, "後述の Step 2 で作ります。今は「作るんだな」と思っておくだけでOKです。")

    add_heading2(doc, "✅ PCの準備")
    add_bullets(doc, [
        "ブラウザは Google Chrome を推奨します。",
        "Chromeで Googleアカウントにログインしておいてください。",
    ])
    add_heading3(doc, "☑ この章で完了したこと")
    add_checklist(doc, [
        "サロン専用 Gmail を用意した",
        "PCに Google Chrome をインストールし、Gmail でログインした",
    ])

    add_page_break(doc)

    # ===== Step 1 =====
    add_step_banner(doc, 1, "Googleスプレッドシートをコピーする",
                    "約10分", "★☆☆（やさしい）", "なし（最初の一歩）")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "このシステムは、Googleスプレッドシート（Google版のエクセル）をデータベース（お客様情報の箱）として使います。開発者から共有されたお手本を自分用にコピーすることで、あなた専用のデータ箱が手に入ります。")

    add_action(doc, "operate", "開発者から届いたメールの「スプレッドシート共有リンク」をクリック。スプレッドシートが開いたら以下を実行してください。")
    add_bullets(doc, [
        "画面左上の「ファイル」メニュー → 「コピーを作成」を選ぶ",
        "「名前」欄に『〇〇サロン_顧客管理』のように分かりやすい名前を入れる",
        "「フォルダ」は「マイドライブ」のままでOK",
        "右下の「コピーを作成」ボタンをクリック",
        "自分用スプレッドシートが新しいタブで開く",
    ], numbered=True)

    add_heading3(doc, "📝 スプレッドシートIDをメモしておく")
    add_para(doc, "ブラウザのアドレスバー（URLが表示されている場所）がこんな形になっています。")
    add_code_block(doc, "https://docs.google.com/spreadsheets/d/【ここの長い文字列がID】/edit")
    add_para(doc, "/d/ と /edit の間の長い文字列（40文字くらい）をメモ帳にコピペして保存。Step 3 で使います。")

    add_action(doc, "check", None)
    add_checklist(doc, [
        "タブのタイトルが「〇〇サロン_顧客管理」になっている",
        "シート下部に『予約台帳』『顧客マスター』『商品マスター』などのタブが並んでいる",
    ])

    add_action(doc, "trouble", None)
    add_bullets(doc, [
        "「権限がありません」と表示される → 開発者に「閲覧権限をください」と連絡",
        "コピーが見当たらない → Googleドライブ（drive.google.com）を検索",
    ])

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, [
        "自分用のスプレッドシートができた",
        "スプレッドシートIDをメモした",
    ])

    add_page_break(doc)

    # ===== Step 2 =====
    add_step_banner(doc, 2, "LINE公式アカウントを作る",
                    "約30分", "★★☆（少し難しい）", "Step 1 完了推奨")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "お客様がLINEから予約できるようにするには、サロン専用のLINEアカウント（LINE公式アカウント）が必要です。無料プランでOK。")

    add_heading3(doc, "A）LINE公式アカウントを作る")
    add_action(doc, "operate", None)
    add_bullets(doc, [
        "ブラウザで https://www.linebiz.com/jp/entry/ を開く",
        "「LINE公式アカウントをはじめる」の青いボタンをクリック",
        "「無料でアカウント開設」を選び、案内に従って進める",
        "サロン名・業種（美容）・連絡先メールを入力 → 完了",
    ], numbered=True)

    add_heading3(doc, "B）Messaging API を有効にする")
    add_para(doc, "お客様との自動やりとり（予約ボット）を動かすには、Messaging API（メッセージを自動で送受信する仕組み）を有効にします。")
    add_action(doc, "operate", None)
    add_bullets(doc, [
        "LINE公式アカウントの管理画面（LINE Official Account Manager）を開く",
        "画面右上の「設定」（歯車アイコン）をクリック",
        "左メニューの「Messaging API」をクリック",
        "「Messaging APIを利用する」ボタンをクリック",
        "LINE Developers というサイトに飛ぶ",
        "「プロバイダー」選択画面で、新規なら「新規プロバイダー作成」→ サロン名を入力",
    ], numbered=True)
    add_callout(doc, "tip",
        "LINE Developers のログイン方法：普段お使いの個人のLINEアカウントでOK。"
        "メールアドレス＋パスワード、または画面に出るQRコードをスマホのLINEカメラで読み取る方法でログインできます。")

    add_callout(doc, "important",
        "「プロバイダー」と「チャネル」の違い — "
        "プロバイダー＝会社・組織の「入れ物」（例：〇〇サロン株式会社）。"
        "チャネル＝そのプロバイダーの中にあるLINE公式アカウント1つ分の設定。"
        "つまり「プロバイダー（会社）の中にチャネル（LINE）が入る」という親子関係です。"
        "サロンが1つなら、プロバイダー1つ＋チャネル1つでOK。")

    add_heading3(doc, "C）チャネルアクセストークン・チャネルシークレットを取得")
    add_para(doc, "この2つの「合言葉（トークン・シークレット）」が、システムとLINEをつなぐカギになります。")
    add_para(doc, "【チャネルシークレットの取得】", bold=True)
    add_bullets(doc, [
        "上部タブの「チャネル基本設定」をクリック",
        "下にスクロールして「チャネルシークレット」を見つける",
        "右の「コピー」ボタン → メモ帳に貼り付け、LINE_CHANNEL_SECRET とラベルを付ける",
    ], numbered=True)
    add_para(doc, "【チャネルアクセストークンの取得】", bold=True)
    add_bullets(doc, [
        "上部タブの「Messaging API設定」をクリック",
        "一番下までスクロール →「チャネルアクセストークン（長期）」欄へ",
        "空欄なら「発行」ボタンをクリック",
        "生成された長い文字列をコピー → LINE_CHANNEL_ACCESS_TOKEN とラベルを付けて保存",
    ], numbered=True)

    add_action(doc, "trouble", None)
    add_bullets(doc, [
        "「Messaging APIを利用する」ボタンが見当たらない → 管理画面の「設定」→「アカウント設定」を試す",
        "トークン発行を押しても何も起きない → ブラウザを再起動・別タブで開き直す",
    ])

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, [
        "LINE公式アカウントが作れた",
        "Messaging API が有効になった",
        "チャネルアクセストークン（長い文字列）をメモした",
        "チャネルシークレット（長い文字列）をメモした",
    ])

    add_page_break(doc)

    # ===== Step 3 =====
    add_step_banner(doc, 3, "GAS（プログラム）を自分用に設定する",
                    "約30分", "★★☆（少し難しい）", "Step 1, 2 完了")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "GAS（ジーエーエス、Google Apps Script の略。Googleが提供する無料のプログラム実行環境）がシステムの本体です。開発者から共有されたGASプロジェクトを、あなたのGoogleアカウントで動くように自分用に複製＆設定します。")

    add_heading3(doc, "A）GASプロジェクトを自分用にコピー")
    add_bullets(doc, [
        "開発者から届いたメールの「GASプロジェクト共有リンク」をクリック",
        "GASエディタ（プログラム画面）が開く",
        "画面左上の「ファイル」→「コピーを作成」を選ぶ（3点メニューの中にある場合も）",
        "自分専用プロジェクトが開く。名前を『〇〇サロン_GAS』に変える",
    ], numbered=True)
    add_callout(doc, "tip",
        "代替手段：開発者が事前に Step 1 のスプレッドシートに紐づくGASを用意してくれている場合、コピー作業は不要です。"
        "スプレッドシートの「拡張機能」→「Apps Script」で開くだけでOK。分からなければ開発者に確認を。")

    add_heading3(doc, "B）スクリプトプロパティを設定する")
    add_para(doc, "スクリプトプロパティ（プログラムに覚えさせる「合言葉」の保管庫）にLINEのトークンなどを登録します。")
    add_para(doc, "GASエディタ左メニューの「⚙ プロジェクトの設定」（歯車）をクリック → 下にスクロール → 「スクリプトプロパティ」→ 「スクリプトプロパティを追加」ボタンで1つずつ追加。")

    add_heading3(doc, "🔑 必須プロパティ（必ず設定）")
    add_table(doc,
        ["プロパティ名", "入れる値", "意味"],
        [
            ["LINE_CHANNEL_ACCESS_TOKEN", "Step 2 の長期トークン", "LINEと通信する合言葉"],
            ["LINE_CHANNEL_SECRET", "Step 2 のシークレット", "LINEの本物確認用"],
            ["SPREADSHEET_ID", "Step 1 のID", "データ箱の場所"],
            ["STORE_EMAIL", "サロンのメールアドレス", "エラー通知の受け取り先"],
            ["ADMIN_PASSWORD", "自分で決めるパスワード（8文字以上推奨）", "管理画面ログイン用"],
        ],
        col_widths=[5.5, 5.5, 5.0])

    add_heading3(doc, "⏳ 後から自動で作られるもの（今は空でOK）")
    add_table(doc,
        ["プロパティ名", "いつ設定されるか"],
        [
            ["OWNER_LINE_ID", "Step 9 のテスト時、ボットに「マイID」と送ると表示される値を登録"],
            ["ADMIN_PASSWORD_HASH", "Step 8 で初めてログイン時に自動生成（平文の ADMIN_PASSWORD は自動削除）"],
            ["BACKUP_FOLDER_ID", "Step 10 の weeklyBackup 初回実行時に自動生成"],
            ["RICH_MENU_ID", "Step 7 の setupRichMenu 実行時に自動生成"],
            ["WEB_API_KEY", "管理画面から最初にAPIアクセスした際に自動生成"],
        ],
        col_widths=[5.5, 10.5])

    add_heading3(doc, "🎨 任意プロパティ（必要なら）")
    add_table(doc,
        ["プロパティ名", "内容"],
        [
            ["MENU_LIST", "メニュー一覧（カンマ区切り）例: カット,カラー,トリートメント"],
            ["SALON_IMAGE_URL", "LINE Flexカード用のサロン画像URL"],
            ["RICHMENU_IMAGE_URL", "リッチメニュー背景画像URL"],
            ["STRIPE_SECRET_KEY", "Stripe（クレカ決済）秘密キー。β版では使いません"],
            ["STRIPE_WEBHOOK_SECRET", "StripeのWebhook署名検証用。β版では使いません"],
            ["EXCHANGE_URL", "ポイント交換用EC画面のURL。β版では使いません"],
        ],
        col_widths=[5.5, 10.5])

    add_callout(doc, "important",
        "β版運用時の注意：STRIPE_SECRET_KEY / STRIPE_WEBHOOK_SECRET / EXCHANGE_URL はEC機能用です。"
        "β版では使わないため設定不要。正式運用で決済機能を有効化する際に開発者と一緒に設定します。")

    add_action(doc, "trouble", None)
    add_bullets(doc, [
        "「スクリプトプロパティ」が見当たらない → 左メニュー一番下の歯車⚙をクリック、中段〜下にあります",
        "値を保存したのに消える → もう一度「保存」ボタン。ブラウザ再読み込みで確認",
    ])

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, [
        "GASプロジェクトが自分のGoogleアカウントに入った",
        "必須5プロパティを全部登録した",
        "自分の ADMIN_PASSWORD を忘れないようメモした",
    ])
    add_callout(doc, "tip",
        "順序の豆知識：基本は Step 3 → Step 4（デプロイ）→ Step 6（initializeSpreadsheet）で進めます。"
        "ただし initializeSpreadsheet は Step 3 直後に実行しても動きます。"
        "その場合は承認画面が1回で済むので少し楽。どちらでもOK。")

    add_page_break(doc)

    # ===== Step 4 =====
    add_step_banner(doc, 4, "ウェブアプリとして公開する（デプロイ）",
                    "約10分", "★★☆（少し難しい）", "Step 3 完了")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "デプロイ（公開設定）を行うと、GASプログラムに専用URLが発行されます。このURLがあって初めて、LINEからシステムに予約データを渡せます。")

    add_action(doc, "operate", None)
    add_bullets(doc, [
        "GASエディタ右上の青い「デプロイ」ボタンをクリック",
        "メニューから「新しいデプロイ」を選択",
        "左上の歯車⚙ →「ウェブアプリ」を選ぶ",
        "「説明」：初回デプロイなど自由",
        "「次のユーザーとして実行」：自分（あなたのメール）",
        "「アクセスできるユーザー」：全員（重要！LINEからアクセスするため）",
        "右下の「デプロイ」ボタンをクリック",
        "「承認が必要です」→「アクセスを承認」→ アカウント選択 →「詳細」→「〇〇（安全ではないページ）に移動」→「許可」",
        "完了画面の「ウェブアプリURL」をコピーしてメモ帳に保存！",
    ], numbered=True)
    add_callout(doc, "warn",
        "「安全ではない」と出ますが、自分のスクリプトなので問題ありません。")

    add_callout(doc, "important",
        "初回と2回目以降のデプロイの違い（重要） — "
        "初回は「新しいデプロイ」。"
        "2回目以降（コード修正時）に「新しいデプロイ」を使うとURLが変わり、LINEのWebhook URLも毎回貼り直しになります。"
        "コード修正反映だけなら、「デプロイを管理」→ 既存デプロイの鉛筆アイコン → バージョンを「新バージョン」に切替 → 「デプロイ」。"
        "これでURLを変えずコードだけ更新できます。")

    add_heading3(doc, "📝 URLの形")
    add_code_block(doc, "https://script.google.com/macros/s/【長いID】/exec")
    add_para(doc, "この【長いID】をデプロイIDと呼びます。Step 5 でLINEに登録します。")

    add_action(doc, "trouble", None)
    add_bullets(doc, [
        "承認画面が繰り返し出る → ブラウザのポップアップブロックを解除",
        "デプロイボタンがグレー → 左のファイル一覧からどれか1つクリックしてから再挑戦",
    ])

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, [
        "ウェブアプリURLがメモできた",
        "アクセスできるユーザーを「全員」にした",
    ])

    add_page_break(doc)

    # ===== Step 5 =====
    add_step_banner(doc, 5, "LINEのWebhook設定",
                    "約10分", "★★☆（少し難しい）", "Step 2, 4 完了")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "Webhook（ウェブフック、LINEにメッセージが届いたらこのURLに通知して！と伝える仕組み）を設定すると、お客様のLINEメッセージをシステムが受け取れます。")

    add_action(doc, "operate", None)
    add_bullets(doc, [
        "LINE Developers のサイトでチャネルの「Messaging API設定」タブを開く",
        "「Webhook設定」の「Webhook URL」欄に、Step 4 のURL（末尾 /exec）を貼り付け",
        "右の「更新」ボタンをクリック",
        "「検証」ボタンをクリック → 「成功」と出ればOK",
        "「Webhookの利用」トグルをONにする",
    ], numbered=True)

    add_callout(doc, "warn",
        "重要：LINE公式アカウント管理画面で自動応答をOFFにすること。"
        "これをやらないとLINEの自動応答がシステムとケンカして予約が動きません。")
    add_bullets(doc, [
        "LINE Official Account Manager を開く",
        "「設定」（歯車）→「応答設定」",
        "「応答メッセージ」→ オフ",
        "「あいさつメッセージ」→ オフ（任意だが推奨）",
        "「Webhook」→ オン",
    ], numbered=True)

    add_action(doc, "trouble", None)
    add_bullets(doc, [
        "Webhook検証で「失敗」→ Step 4 の「アクセスできるユーザー: 全員」を確認",
        "URLが違う → 末尾が /exec で終わっているか確認",
    ])

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, [
        "Webhook URLが貼れた・検証成功",
        "応答メッセージをOFFにした",
        "Webhook利用をONにした",
    ])

    add_page_break(doc)

    # ===== Step 6 =====
    add_step_banner(doc, 6, "初期化スクリプトを実行する",
                    "約5分", "★☆☆（やさしい）", "Step 3, 4 完了")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "スプレッドシートに、予約を記録する表や顧客を記録する表など、必要なシートを自動で全部作るためです。")

    add_action(doc, "operate", None)
    add_bullets(doc, [
        "GASエディタの関数ドロップダウン（「関数を選択」）をクリック",
        "initializeSpreadsheet を選択",
        "「▶ 実行」ボタンをクリック",
        "「承認が必要です」→ Step 4 と同じ手順で許可",
        "「スプレッドシートの初期化が完了しました」アラートが出ればOK",
    ], numbered=True)

    add_action(doc, "check", None)
    add_bullets(doc, [
        "スプレッドシートのタブが 予約台帳／顧客マスター／商品マスター／注文履歴／来店記録／支払い記録／交換申請 の 7つ に増えている",
    ])

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, [
        "初期化スクリプトが正常終了した",
        "スプレッドシートに必要なシートが揃っている",
    ])

    add_page_break(doc)

    # ===== Step 7 =====
    add_step_banner(doc, 7, "リッチメニューを設定する",
                    "約5分", "★☆☆（やさしい）", "Step 3, 4 完了")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "リッチメニュー（LINEトーク画面下の大きなボタン群）を設定すると、お客様が「予約」「予約確認」「お問い合わせ」をタップ1つでできます。")

    add_action(doc, "operate", None)
    add_bullets(doc, [
        "関数ドロップダウンから setupRichMenu を選択",
        "「▶ 実行」をクリック",
        "ログに「リッチメニュー作成完了」と出れば成功",
        "RICH_MENU_ID が自動でスクリプトプロパティに保存される",
    ], numbered=True)

    add_action(doc, "check", "LINEでサロン公式アカウントを友だち追加（Step 9 でOK）した後、トーク画面下にメニューボタンが出ていればOK")
    add_action(doc, "trouble", None)
    add_bullets(doc, [
        "エラーが出る → LINE_CHANNEL_ACCESS_TOKEN が正しいか確認",
        "メニューが表示されない → LINEアプリ再起動、またはトーク画面を閉じて開き直す",
    ])

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, ["setupRichMenu が正常終了した"])

    add_page_break(doc)

    # ===== Step 8 =====
    add_step_banner(doc, 8, "管理画面にログインしてみる",
                    "約10分", "★☆☆（やさしい）", "Step 3 完了")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "ブラウザから予約一覧・顧客一覧・売上を見るための管理画面にログインできるか確認します。")

    add_callout(doc, "important",
        "管理画面URLの運用について — "
        "パターンA（共有型）：開発者が1つのURLを運用し、GAS_URLだけサロンごとに差し替え（最も手軽）。"
        "パターンB（複製型）：サロンごとにGitHubをfork（コピー）し、各サロン専用URLを持つ。"
        "どちらか分からない場合は開発者に「管理画面の運用形態はA・Bどちらですか？」と聞いてください。")

    add_action(doc, "operate", None)
    add_bullets(doc, [
        "管理画面URLをブラウザで開く（例: https://xxxx.github.io/salon-repo/）",
        "パスワード欄に Step 3 の ADMIN_PASSWORD を入力",
        "「ログイン状態を保持する」は専用PCならON、共有PCならOFF推奨",
        "「ログイン」→ ダッシュボードが表示されれば成功",
    ], numbered=True)

    add_heading3(doc, "🎨 自分好みの文言を入れる（運用設定タブ）")
    add_bullets(doc, [
        "左メニューから「設定」をクリック",
        "「運用設定」タブを開く",
        "キャンセルポリシー文言を自分のサロン用に書き換える",
        "「保存」をクリック",
    ], numbered=True)
    add_code_block(doc,
        "※ キャンセルは前日までにお願いします。当日キャンセルは50%、無断キャンセルは100%のキャンセル料がかかります。")

    add_action(doc, "trouble", None)
    add_bullets(doc, [
        "「パスワードが違います」→ Step 3 でメモしたパスワードを確認",
        "画面が真っ白 → F5で再読み込み。config.js の GAS_URL が正しいか開発者に確認",
    ])

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, [
        "管理画面にログインできた",
        "キャンセルポリシーを自分のサロン仕様に書き換えた",
    ])

    add_page_break(doc)

    # ===== Step 9 =====
    add_step_banner(doc, 9, "テスト予約をしてみる",
                    "約15分", "★★☆（少し難しい）", "Step 1〜8 完了")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "ここまでの設定がすべて正しく動くか、自分のスマホで実際に予約して確認します。")

    add_action(doc, "operate", None)
    add_bullets(doc, [
        "スマホのLINEアプリを開く",
        "LINE Official Account Manager の「友だちを増やす」からQRコードを取得",
        "スマホでサロン公式アカウントを友だち追加",
        "トーク画面を開き、まず「マイID」と送信",
        "U から始まる長い文字列（あなたのLINE ID）が返ってくる → コピーしてメモ",
        "GASエディタのスクリプトプロパティ OWNER_LINE_ID にこの値を登録",
        "再度LINEでリッチメニューから「予約」をタップ",
        "週間カレンダーが表示されたら好きな○の時間帯をタップ",
        "電話番号→（新規なら）名前→メニュー選択→確認画面→「予約確定」",
        "「予約が完了しました」と返れば大成功！🎉",
    ], numbered=True)

    add_action(doc, "check", None)
    add_bullets(doc, [
        "管理画面の「予約管理」にテスト予約が追加されている",
        "「顧客一覧」にあなたの名前がある",
        "OWNER_LINE_ID 設定後、オーナー通知LINEが届く",
    ])

    add_action(doc, "trouble", None)
    add_bullets(doc, [
        "カレンダーが出ない → Webhook設定（Step 5）を再確認",
        "「このメッセージは対応していません」→ 応答メッセージがONのまま。Step 5 で OFF に",
    ])

    add_heading3(doc, "🩺 「マイID」と送っても何も返ってこない場合の切り分け")
    add_checklist(doc, [
        "①Webhook利用がON？ — LINE Developers → Messaging API設定 → 「Webhookの利用」トグル",
        "②応答メッセージがOFF？ — LINE Official Account Manager → 設定 → 応答設定",
        "③Webhook URLは正しい？ — 末尾 /exec、「検証」で成功が出るか",
        "④友だち追加できている？ — 友だち一覧に入っているか、ブロックしていないか",
    ])

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, [
        "友だち追加できた",
        "OWNER_LINE_ID を設定した",
        "テスト予約が最後まで通った",
    ])

    add_page_break(doc)

    # ===== Step 10 =====
    add_step_banner(doc, 10, "定期トリガーをセットする",
                    "約5分", "★☆☆（やさしい）", "Step 3 完了")
    add_heading3(doc, "🤔 なぜやるの？")
    add_para(doc, "トリガー（決まった時間に自動でプログラムを動かす仕組み）をセットすると、以下が自動実行されます。")
    add_bullets(doc, [
        "毎日20時 → 翌日予約のリマインドLINE自動送信",
        "毎週日曜2時 → スプレッドシート全体の自動バックアップ",
        "毎月1日3時 → LINEログの整理（古いログを別シートに退避）",
    ])

    add_action(doc, "operate", None)
    add_bullets(doc, [
        "関数ドロップダウンから installTriggers を選択",
        "「▶ 実行」をクリック",
        "承認が出たら許可",
        "「トリガーの設定が完了しました」と表示されればOK",
    ], numbered=True)

    add_action(doc, "check", "GASエディタ左メニューの「⏰ トリガー」（時計アイコン）をクリックすると、3つのトリガーが登録されている")

    add_heading3(doc, "☑ このステップで完了したこと")
    add_checklist(doc, ["3つのトリガーが登録された"])

    add_page_break(doc)

    # ===== 完了メッセージ =====
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🎊 おめでとうございます！セットアップ完了！")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = COLOR_H1
    apply_jp_font(run)
    add_blank(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "お疲れ様でした。これで、あなたのサロン用システムが24時間稼働状態になりました。", size=12)

    add_page_break(doc)

    # ===== 日常運用ガイド =====
    add_heading1(doc, "日常運用ガイド")
    add_heading2(doc, "🌅 毎日のチェック（所要1〜2分）")
    add_bullets(doc, [
        "朝、管理画面の「ダッシュボード」を開く",
        "「本日の予約」を確認",
        "🚨「コース残回数アラート」に名前があれば、次回来店時に声かけ",
    ], numbered=True)
    add_heading2(doc, "📅 週1回のチェック（日曜夜など）")
    add_bullets(doc, [
        "ダッシュボード上部の「💾 最新バックアップ」カードが緑色（3日以内）か確認",
        "オレンジなら手動で weeklyBackup を実行（FAQ参照）",
    ], numbered=True)
    add_heading2(doc, "📊 月1回のチェック")
    add_bullets(doc, [
        "管理画面の「売上レポート」で月次売上を確認",
        "顧客マスターで最終来店日が2ヶ月以上経っているお客様にフォローLINEを送る",
    ], numbered=True)

    add_page_break(doc)

    # ===== FAQ =====
    add_heading1(doc, "こんなとき どうする（FAQ）")

    add_heading2(doc, "❓ LINEからの予約が届かない")
    add_checklist(doc, [
        "Step 5 の Webhook URLが正しく貼られているか？",
        "Webhook検証で「成功」になっているか？",
        "Webhookの利用がONか？",
        "応答メッセージがOFFか？",
        "Messaging API設定 → 応答メッセージもOFFか？",
    ])

    add_heading2(doc, "❓ 管理画面にログインできない（パスワードを忘れた）")
    add_bullets(doc, [
        "GASエディタを開く",
        "⚙ プロジェクトの設定 → スクリプトプロパティ",
        "ADMIN_PASSWORD_HASH の行を削除",
        "ADMIN_PASSWORD に新しいパスワードを入れる",
        "次回ログイン時に自動で再ハッシュ化",
    ], numbered=True)

    add_heading2(doc, "❓ バックアップが警告色（オレンジ）になっている")
    add_callout(doc, "tip",
        "前提：BACKUP_FOLDER_ID は weeklyBackup を初めて実行したときにDrive内に自動生成されます。"
        "installTriggers 実行だけではフォルダは作られません。"
        "セットアップ直後にオレンジは正常。すぐ緑にしたい場合は手動実行してください。")
    add_bullets(doc, [
        "GASエディタを開く",
        "関数ドロップダウンから weeklyBackup を選択",
        "「▶ 実行」をクリック",
        "初回は承認画面→許可。Driveに「サロン顧客管理_バックアップ」フォルダが自動作成",
        "管理画面ダッシュボードを再読み込みし緑になればOK",
    ], numbered=True)

    add_heading2(doc, "❓ スプレッドシートの列を間違えて消した／おかしくなった")
    add_bullets(doc, [
        "Googleドライブを開く",
        "マイドライブ →「サロン顧客管理_バックアップ」フォルダ",
        "最新日付のバックアップを右クリック →「コピーを作成」",
        "コピーを開いて必要な行をコピー → 元のスプレッドシートに貼り付け",
        "全部ダメな時は、そのバックアップを新しい本番に（SPREADSHEET_ID 更新）",
    ], numbered=True)

    add_heading2(doc, "❓ 顧客が「パスワードを忘れた」と言ってきた")
    add_bullets(doc, [
        "管理画面にログイン",
        "「顧客一覧」からその方の名前をクリック",
        "「🔑 ログインパスワード」カードの「パスワードをリセット」オレンジボタン",
        "確認ダイアログで「OK」→ リセット完了",
        "お客様にLINEで「パスワードをリセットしました。もう一度ご登録ください」と連絡",
    ], numbered=True)

    add_heading2(doc, "❓ 予約のキャンセル通知がオーナーに届かない")
    add_bullets(doc, ["スクリプトプロパティ OWNER_LINE_ID が正しく設定されているか確認（Step 9 参照）"])

    add_heading2(doc, "❓ 新しいスタッフを追加したい")
    add_bullets(doc, ["現在オーナー1人＋お客様構成。スタッフ個別ログインは未実装。開発者に相談を。"])

    add_heading2(doc, "❓ Webhook検証で401/403エラー")
    add_bullets(doc, [
        "GASエディタ右上の「デプロイ」→「デプロイを管理」",
        "現在のデプロイの鉛筆アイコンをクリック",
        "「アクセスできるユーザー」が「全員」か確認",
        "「全員」に変更 →「デプロイ」でURLを保ったまま更新",
        "LINE Developers 側で「検証」を再実行",
    ], numbered=True)

    add_heading2(doc, "❓ SPREADSHEET_ID を間違えた")
    add_bullets(doc, [
        "正しいスプレッドシートのIDをコピー（Step 1 参照）",
        "GASエディタ → ⚙ → スクリプトプロパティ",
        "SPREADSHEET_ID の値を書き換えて保存",
        "間違ったシートに書き込まれたデータがあれば手動でコピペ",
        "念のため initializeSpreadsheet を再実行",
    ], numbered=True)

    add_heading2(doc, "❓ LINEの月200通無料枠を超えそう／超えた")
    add_bullets(doc, [
        "状況確認：LINE Official Account Manager →「分析」→「メッセージ通数」",
        "対策1（節約）：リマインド内容を簡潔に／不要な一斉送信を減らす",
        "対策2（プラン変更）：ライトプラン（月5,000円で5,000通）などにアップグレード",
        "対策3：installTriggers で登録された dailyReminder を⏰トリガー画面から削除",
    ])

    add_page_break(doc)

    # ===== 困ったら連絡先 =====
    add_heading1(doc, "困ったら連絡先")
    add_heading2(doc, "📧 開発者への問い合わせテンプレ")
    add_para(doc, "解決しない問題があれば、以下のテンプレをコピーして開発者にメールまたはLINEしてください。状況を具体的に書くと早く解決します。")
    add_code_block(doc,
        "【サロン名】〇〇サロン\n"
        "\n"
        "【発生日時】2026年◯月◯日 ◯時ごろ\n"
        "\n"
        "【何をしたか】\n"
        "（例：LINEで「予約」と送った／管理画面の「完了」ボタンを押した）\n"
        "\n"
        "【どうなったか】\n"
        "（例：「このメッセージは対応していません」と返ってきた／画面が真っ白になった）\n"
        "\n"
        "【スクリーンショット】\n"
        "（あれば添付。スマホなら電源ボタン+音量下で撮影）\n"
        "\n"
        "【試したこと】\n"
        "（例：FAQの「LINEからの予約が届かない」を確認したが改善せず）"
    )

    add_heading2(doc, "💡 アドバイス")
    add_callout(doc, "important",
        "焦らないこと。システムは自動バックアップがあるので、データが完全に消えることはほぼありません。")
    add_callout(doc, "warn",
        "勝手にコードをいじらないこと。GASエディタでうっかりファイルを変更・削除すると復旧が大変です。")
    add_callout(doc, "warn",
        "スクリプトプロパティの LINE_CHANNEL_ACCESS_TOKEN などを絶対に人に見せないこと。"
        "漏れると他人にLINEを乗っ取られる可能性があります。")

    add_page_break(doc)

    # ===== 用語集 =====
    add_heading1(doc, "📚 用語集")
    add_para(doc, "初心者向けに、このマニュアルで登場する専門用語をやさしく解説します。")
    add_table(doc,
        ["用語", "初心者向けの意味"],
        [
            ["GAS", "Google Apps Script。Googleが提供する無料のプログラム実行環境。このシステムの本体。"],
            ["Webhook", "「ある出来事が起きたらこのURLに通知して！」とお願いする仕組み。LINE→システムの通知経路。"],
            ["デプロイ", "作ったプログラムを「公開状態」にして、他から使えるURLを発行すること。"],
            ["チャネル", "LINE公式アカウント1つ分の設定。プロバイダー（会社）の中にチャネル（LINE）が入る。"],
            ["プロバイダー", "LINE Developers 上の「会社・組織」の入れ物。中に複数のチャネルを作れる。"],
            ["リッチメニュー", "LINEトーク画面下に出る大きなボタン群。タップ1つで予約などの機能を呼び出せる。"],
            ["Flexカード", "LINEに画像・ボタン付きで送れる装飾メッセージ。予約確認などで使用。"],
            ["トリガー", "GASで「決まった時間に自動で関数を動かす」スケジューラーの仕組み。"],
            ["スクリプトプロパティ", "GASプロジェクトに覚えさせる「合言葉」の保管庫。トークンや設定値を安全に保存。"],
            ["ハッシュ化", "パスワードを元に戻せない形に変換する処理。漏洩時の被害を減らせる。"],
            ["デプロイID", "デプロイ完了時に発行される識別子。ウェブアプリURLの /s/〜/exec の間の長い文字列。"],
            ["スプレッドシートID", "スプレッドシートのURL /d/〜/edit の間の長い文字列。データ箱の場所を示す。"],
        ],
        col_widths=[4.0, 12.0])

    # 保存
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
